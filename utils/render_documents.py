#!/usr/bin/env python3
"""
render_documents.py  — S3: Chuyển .md  →  .docx  và  .xlsx (bảng)

Cách dùng:
    python utils/render_documents.py                      # Toàn bộ Outputs/ của 5 Agent
    python utils/render_documents.py --agent AG-05        # Chỉ AG-05
    python utils/render_documents.py --input file.md      # File đơn lẻ
    python utils/render_documents.py --input f.md --output out.docx
"""

import os
import re
import argparse
from pathlib import Path

# ── Kiểm tra dependencies theo cách graceful ──────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt
    DOCX_OK = True
except ImportError:
    DOCX_OK = False
    print("⚠️  python-docx chưa cài. Chạy: pip install -r requirements.txt")

try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill
    XLSX_OK = True
except ImportError:
    XLSX_OK = False
    print("⚠️  openpyxl chưa cài. Chạy: pip install -r requirements.txt")


# ==========================================
# REGEX PATTERNS
# ==========================================
HEADING_RE    = re.compile(r"^(#{1,6})\s+(.+)$")
TABLE_ROW_RE  = re.compile(r"^\|(.+)\|$")
TABLE_SEP_RE  = re.compile(r"^\|[\s\-|:]+\|$")
CODE_FENCE_RE = re.compile(r"^```")
HR_RE         = re.compile(r"^[-*_]{3,}$")
UL_RE         = re.compile(r"^[ \t]*[-*+]\s+(.+)$")
OL_RE         = re.compile(r"^[ \t]*\d+\.\s+(.+)$")
INLINE_RE     = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


# ==========================================
# MARKDOWN → DOCX
# ==========================================
def _add_run_formatted(para, text: str):
    """Thêm run vào paragraph, xử lý **bold**, *italic*, `code` inline."""
    for seg in INLINE_RE.split(text):
        if seg.startswith("**") and seg.endswith("**"):
            run = para.add_run(seg[2:-2])
            run.bold = True
        elif seg.startswith("*") and seg.endswith("*") and len(seg) > 2:
            run = para.add_run(seg[1:-1])
            run.italic = True
        elif seg.startswith("`") and seg.endswith("`"):
            run = para.add_run(seg[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif seg:
            para.add_run(seg)


def _flush_table(doc, table_rows: list):
    """Render bảng Markdown tích lũy vào Document."""
    if not table_rows:
        return
    data_rows = [r for r in table_rows if not TABLE_SEP_RE.match(r)]
    if not data_rows:
        return

    parsed = []
    for row in data_rows:
        cells = [c.strip() for c in row.strip("|").split("|")]
        parsed.append(cells)

    col_count = max(len(r) for r in parsed)
    t = doc.add_table(rows=len(parsed), cols=col_count)
    t.style = "Table Grid"
    for ri, row in enumerate(parsed):
        for ci in range(col_count):
            val = row[ci] if ci < len(row) else ""
            cell = t.cell(ri, ci)
            cell.text = val
            if ri == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph()


def md_to_docx(md_text: str, output_path: str):
    """Chuyển nội dung Markdown thành file .docx."""
    if not DOCX_OK:
        print(f"  ⚠️  Bỏ qua {output_path} — python-docx chưa cài.")
        return

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_lines: list = []
    table_rows: list = []

    while i < len(lines):
        line = lines[i]

        # Code block fence
        if CODE_FENCE_RE.match(line):
            if not in_code:
                _flush_table(doc, table_rows); table_rows = []
                in_code = True
                code_lines = []
            else:
                in_code = False
                para = doc.add_paragraph()
                for cl in code_lines:
                    run = para.add_run(cl + "\n")
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Table row
        if TABLE_ROW_RE.match(line):
            table_rows.append(line)
            i += 1
            continue
        else:
            _flush_table(doc, table_rows)
            table_rows = []

        # Heading
        m = HEADING_RE.match(line)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(m.group(2), level=level)
            i += 1
            continue

        # Horizontal rule
        if HR_RE.match(line) and line.strip():
            doc.add_paragraph().add_run("─" * 50)
            i += 1
            continue

        # Unordered list
        ul = UL_RE.match(line)
        if ul:
            para = doc.add_paragraph(style="List Bullet")
            _add_run_formatted(para, ul.group(1))
            i += 1
            continue

        # Ordered list
        ol = OL_RE.match(line)
        if ol:
            para = doc.add_paragraph(style="List Number")
            _add_run_formatted(para, ol.group(1))
            i += 1
            continue

        # Empty line
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        # Normal paragraph
        para = doc.add_paragraph()
        _add_run_formatted(para, line)
        i += 1

    _flush_table(doc, table_rows)
    doc.save(output_path)
    print(f"  📄 Đã tạo: {output_path}")


# ==========================================
# MARKDOWN TABLES → XLSX
# ==========================================
def extract_tables(md_text: str):
    """Trích tất cả bảng từ Markdown. Trả về list of list-of-rows."""
    tables, current = [], []
    for line in md_text.splitlines():
        if TABLE_ROW_RE.match(line) and not TABLE_SEP_RE.match(line):
            cells = [c.strip() for c in line.strip("|").split("|")]
            current.append(cells)
        else:
            if current:
                tables.append(current)
                current = []
    if current:
        tables.append(current)
    return tables


def tables_to_xlsx(tables: list, output_path: str):
    """Ghi danh sách bảng vào .xlsx (mỗi bảng 1 sheet)."""
    if not XLSX_OK:
        print(f"  ⚠️  Bỏ qua {output_path} — openpyxl chưa cài.")
        return
    if not tables:
        return

    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    hdr_fill = PatternFill("solid", fgColor="2D3748")
    hdr_font = Font(bold=True, color="FFFFFF", name="Calibri", size=11)
    nrm_font = Font(name="Calibri", size=10)
    center   = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for idx, table in enumerate(tables):
        ws = wb.create_sheet(title=f"Table_{idx + 1}")
        for ri, row in enumerate(table):
            for ci, value in enumerate(row, start=1):
                cell = ws.cell(row=ri + 1, column=ci, value=value)
                cell.font      = hdr_font if ri == 0 else nrm_font
                cell.alignment = center
                if ri == 0:
                    cell.fill = hdr_fill
        # Auto column width
        for col in ws.columns:
            width = max((len(str(c.value or "")) for c in col), default=8)
            ws.column_dimensions[col[0].column_letter].width = min(width + 4, 50)

    wb.save(output_path)
    print(f"  📊 Đã tạo: {output_path}  ({len(tables)} bảng)")


# ==========================================
# BATCH: quét toàn bộ Outputs/
# ==========================================
AGENT_DIRS = {
    "AG-01": "01_Analyst",
    "AG-02": "02_Architect",
    "AG-03": "03_Developer",
    "AG-04": "04_QA_Reviewer",
    "AG-05": "05_Final_Doc",
}


def render_all(base_dir: str, agent_filter: str = None):
    """Quét Outputs/ của mỗi Agent, render .md → .docx (& .xlsx nếu có bảng)."""
    total = 0
    for agent_id, folder in AGENT_DIRS.items():
        if agent_filter and agent_filter != agent_id:
            continue
        outputs_dir = os.path.join(base_dir, folder, "Outputs")
        if not os.path.isdir(outputs_dir):
            continue
        for root, _, files in os.walk(outputs_dir):
            for fname in files:
                if not fname.endswith(".md"):
                    continue
                md_path = os.path.join(root, fname)
                stem = Path(fname).stem
                with open(md_path, encoding="utf-8") as f:
                    content = f.read()

                # → .docx
                md_to_docx(content, os.path.join(root, stem + ".docx"))

                # → .xlsx (chỉ khi có bảng)
                tbls = extract_tables(content)
                if tbls:
                    tables_to_xlsx(tbls, os.path.join(root, stem + ".xlsx"))

                total += 1

    print(f"\n✅ Render xong {total} file(s).")


# ==========================================
# ENTRY POINT
# ==========================================
if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Render .md → .docx / .xlsx từ Outputs của các Agent",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""Ví dụ:
  python utils/render_documents.py                      # Toàn bộ Outputs
  python utils/render_documents.py --agent AG-05        # Chỉ AG-05
  python utils/render_documents.py --input report.md    # File đơn
""",
    )
    parser.add_argument("--agent",  help="Chỉ render agent này (VD: AG-05)")
    parser.add_argument("--input",  help="Render 1 file .md cụ thể")
    parser.add_argument("--output", help="Đường dẫn output .docx (dùng với --input)")
    args = parser.parse_args()

    _utils_dir = os.path.dirname(os.path.abspath(__file__))
    _base_dir  = os.path.dirname(_utils_dir)

    if args.input:
        if not os.path.exists(args.input):
            print(f"❌ File không tồn tại: {args.input}")
            raise SystemExit(1)
        with open(args.input, encoding="utf-8") as f:
            _content = f.read()
        _out = args.output or str(Path(args.input).with_suffix(".docx"))
        print(f"📄 Render: {args.input} → {_out}")
        md_to_docx(_content, _out)
        _tbls = extract_tables(_content)
        if _tbls:
            _xlsx_out = str(Path(_out).with_suffix(".xlsx"))
            tables_to_xlsx(_tbls, _xlsx_out)
    else:
        render_all(_base_dir, agent_filter=args.agent)
